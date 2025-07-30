# streamlit_app.py
import streamlit as st
import pandas as pd
from modul_rotation import lade_daten, finde_aktuellen_bereich
from openpyxl import load_workbook
from datetime import datetime
import os

EXCEL_PATH = r"C:\Users\ASYAKKA\Mercedes-Benz (corpdir.onmicrosoft.com)\DWT_UTeam Werk 10 - General\08_Rotation UTeam\Projekt_UTeam_Digitalisierung\Masterliste_UTeam.xlsx"
OUTPUT_DIR  = r"C:\Users\ASYAKKA\Mercedes-Benz (corpdir.onmicrosoft.com)\DWT_UTeam Werk 10 - General\Vorstellung_Asya_Test"

# Seiten‑Navigation
st.sidebar.title("UTeam Rotations-Tool")
page = st.sidebar.radio("", [
    "🏠 Startseite",
    "👥 Übersicht",
    "🔄 Rotationsplan",
    "📝 Vorstellungsgespräch"
])

# --- Startseite ---
if page == "🏠 Startseite":
    st.title("UTeam Rotations-Tool")
    st.write("Wähle links einen Bereich aus.")

# --- Mitarbeiterübersicht ---
elif page == "👥 Übersicht":
    st.title("Mitarbeiterübersicht")
    df = lade_daten()  # dein bestehender Loader
    if df is None:
        st.error("Excel konnte nicht geladen werden.")
    else:
        df["Aktueller Bereich"] = df.apply(finde_aktuellen_bereich, axis=1)
        df["Aktuelles Austrittsdatum"] = (
            pd.to_datetime(df["Aktuelles Austrittsdatum"], errors="coerce")
              .dt.strftime("%d.%m.%Y")
        )
        st.dataframe(df[["Vorname","Nachname","Aktueller Bereich","Aktuelles Austrittsdatum"]])

# --- Rotationsplan ---
elif page == "🔄 Rotationsplan":
    st.title("Rotationsplan")
    df = lade_daten(sheet_name="Rotationsplan")
    if df is None:
        st.error("Rotationsplan-Tab konnte nicht geladen werden.")
    else:
        # deine Logik zum Aufbereiten
        st.dataframe(df)  # oder st.table, je nach Wunsch

# --- Vorstellungsgespräch anlegen ---
elif page == "📝 Vorstellungsgespräch":
    st.title("Vorstellungsgespräch anlegen")
    with st.form("form_vorstellung"):
        col1, col2 = st.columns(2)
        with col1:
            vor  = st.text_input("Vorname *")
            nach = st.text_input("Nachname *")
            geb  = st.text_input("Geburtsdatum * (DD.MM.YYYY)")
            eins = st.text_input("Aktueller Einsatz *")
            kst  = st.text_input("Stamm-Kostenstelle *", value="010-")
            gender = st.selectbox("Geschlecht", ["m","w","d"])
            fork   = st.selectbox("Staplerschein", ["ja","nein","k.A."])
        with col2:
            lauf = st.text_area("Laufbahn")
            qual = st.text_area("Qualifikation")
            wunsch = st.text_area("Wunsch")
            sonst = st.text_area("Sonstiges")
        submitted = st.form_submit_button("Gespräch anlegen")
        if submitted:
            # 1) Word erzeugen
            os.makedirs(OUTPUT_DIR, exist_ok=True)
            date_str = datetime.today().strftime("%d.%m.%Y")
            fname = f"{nach}_{vor}_{date_str.replace('.','-')}.docx"
            fullpath = os.path.join(OUTPUT_DIR, fname)
            from docx import Document
            doc = Document()
            doc.add_heading("Vorstellungsgespräch", level=1)
            doc.add_paragraph(f"Datum: {date_str}")
            tbl = doc.add_table(rows=0, cols=2)
            tbl.style = "Light List Accent 1"
            def add_row(label, val, bold=True):
                cells = tbl.add_row().cells
                run = cells[0].paragraphs[0].add_run(label)
                if bold: run.bold = True
                cells[1].text = val
            for lbl,val,b in [
                ("Vorname:", vor, True),
                ("Nachname:", nach, True),
                ("Geburtsdatum:", geb, True),
                ("Aktueller Einsatz:", eins, True),
                ("Stamm-Kostenstelle:", kst, True),
                ("Geschlecht:", gender, True),
                ("Staplerschein:", fork, True),
                ("Laufbahn:", lauf, False),
                ("Qualifikation:", qual, False),
                ("Wunsch:", wunsch, False),
                ("Sonstiges:", sonst, False),
            ]:
                add_row(lbl, str(val), b)
            doc.save(fullpath)

            # 2) Excel ergänzen
            from openpyxl import load_workbook
            wb = load_workbook(EXCEL_PATH)
            ws = wb["Masterlist"]
            # Spalten A=Vorname, B=Nachname, D=Geburtsdatum
            ws.append([vor, nach, None, geb])
            wb.save(EXCEL_PATH)

            st.success(f"Word: {fullpath}\nEintrag in Excel hinzugefügt.")

